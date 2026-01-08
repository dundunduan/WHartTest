"""
çŸ¥è¯†åº“æœåŠ¡æ¨¡å—
æä¾›æ–‡æ¡£å¤„ç†ã€å‘é‡åŒ–ã€æ£€ç´¢ç­‰æ ¸å¿ƒåŠŸèƒ½
"""
import os
import time
import hashlib
from typing import List, Dict, Any
import nltk
from django.conf import settings

# --- NLTK æ•°æ®è·¯å¾„é…ç½® ---
# å°†é¡¹ç›®å†…éƒ¨çš„ 'nltk_data' ç›®å½•æ·»åŠ åˆ° NLTK çš„æœç´¢è·¯å¾„ä¸­
# è¿™ä½¿å¾—é¡¹ç›®åœ¨ä»»ä½•ç¯å¢ƒä¸­éƒ½èƒ½æ‰¾åˆ°å¿…è¦çš„æ•°æ®ï¼Œæ— éœ€ç³»ç»Ÿçº§å®‰è£…
LOCAL_NLTK_DATA_PATH = os.path.join(settings.BASE_DIR, 'nltk_data')
if os.path.exists(LOCAL_NLTK_DATA_PATH):
    if LOCAL_NLTK_DATA_PATH not in nltk.data.path:
        nltk.data.path.insert(0, LOCAL_NLTK_DATA_PATH)
        print(f"NLTK data path prepended with: {LOCAL_NLTK_DATA_PATH}")

# è®¾ç½®å®Œå…¨ç¦»çº¿æ¨¡å¼ï¼Œé¿å…ä»»ä½•ç½‘ç»œè¯·æ±‚
os.environ['TRANSFORMERS_OFFLINE'] = '1'
os.environ['HF_DATASETS_OFFLINE'] = '1'
os.environ['HF_HUB_OFFLINE'] = '1'
os.environ['TOKENIZERS_PARALLELISM'] = 'false'
# ç¦ç”¨ç½‘ç»œè¿æ¥
os.environ['HF_HUB_DISABLE_TELEMETRY'] = '1'
os.environ['HF_HUB_DISABLE_PROGRESS_BARS'] = '1'
os.environ['HF_HUB_DISABLE_SYMLINKS_WARNING'] = '1'
# è®¾ç½®æçŸ­çš„è¿æ¥è¶…æ—¶ï¼Œå¼ºåˆ¶å¿«é€Ÿå¤±è´¥
os.environ['HF_HUB_TIMEOUT'] = '1'
os.environ['REQUESTS_TIMEOUT'] = '1'
from django.conf import settings
from django.utils import timezone
from langchain_community.document_loaders import (
    PyPDFLoader, Docx2txtLoader, UnstructuredPowerPointLoader,
    TextLoader, UnstructuredMarkdownLoader, UnstructuredHTMLLoader,
    WebBaseLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    VectorParams,
    PointStruct,
    SparseVector,
    SparseVectorParams,
    SparseIndexParams,
    NamedVector,
    NamedSparseVector,
    models,
)
from langchain_core.documents import Document as LangChainDocument
from .models import KnowledgeBase, Document, DocumentChunk, QueryLog, KnowledgeGlobalConfig
import logging
import requests
import uuid
from typing import List, Optional, Dict
from langchain.embeddings.base import Embeddings

# å°è¯•å¯¼å…¥ FastEmbed ç”¨äº BM25 ç¨€ç–ç¼–ç 
# æ³¨æ„ï¼šéœ€è¦åœ¨å¯¼å…¥å‰ä¸´æ—¶ç¦ç”¨ç¦»çº¿æ¨¡å¼
FASTEMBED_AVAILABLE = False
SparseTextEmbedding = None

def _init_fastembed():
    """å»¶è¿Ÿåˆå§‹åŒ– FastEmbedï¼ˆé¿å…æ¨¡å—çº§åˆ«çš„ç¦»çº¿æ¨¡å¼å½±å“ï¼‰"""
    global FASTEMBED_AVAILABLE, SparseTextEmbedding
    if FASTEMBED_AVAILABLE:
        return True
    
    # ä¸´æ—¶ç¦ç”¨ç¦»çº¿æ¨¡å¼
    offline_vars = ['HF_HUB_OFFLINE', 'TRANSFORMERS_OFFLINE', 'HF_DATASETS_OFFLINE']
    old_values = {var: os.environ.pop(var, None) for var in offline_vars}
    
    try:
        from fastembed import SparseTextEmbedding as _SparseTextEmbedding
        SparseTextEmbedding = _SparseTextEmbedding
        FASTEMBED_AVAILABLE = True
        return True
    except ImportError:
        return False
    finally:
        # æ¢å¤ç¯å¢ƒå˜é‡
        for var, val in old_values.items():
            if val is not None:
                os.environ[var] = val

logger = logging.getLogger(__name__)


class SparseBM25Encoder:
    """åŸºäº FastEmbed çš„ BM25 ç¨€ç–ç¼–ç å™¨"""

    DEFAULT_MODEL = "Qdrant/bm25"

    def __init__(self, model_name: Optional[str] = None):
        # åˆå§‹åŒ– FastEmbedï¼ˆå»¶è¿Ÿå¯¼å…¥ï¼‰
        if not _init_fastembed():
            raise ImportError("éœ€è¦å®‰è£… fastembed æ‰èƒ½å¯ç”¨ BM25 ç¨€ç–å‘é‡: pip install fastembed")
        
        self.model_name = model_name or self.DEFAULT_MODEL
        
        # æ£€æŸ¥æ˜¯å¦å­˜åœ¨æœ¬åœ°ç¼“å­˜ï¼ˆDocker éƒ¨ç½²æ—¶æ¨¡å‹å·²é¢„ä¸‹è½½ï¼‰
        cache_path = os.environ.get('FASTEMBED_CACHE_PATH', os.path.expanduser('~/.cache/fastembed'))
        model_cache_exists = os.path.isdir(cache_path) and any(
            'bm25' in d.lower() for d in os.listdir(cache_path)
        ) if os.path.exists(cache_path) else False
        
        if model_cache_exists:
            # æœ‰æœ¬åœ°ç¼“å­˜æ—¶ï¼Œä¿æŒç¦»çº¿æ¨¡å¼ï¼Œç›´æ¥åŠ è½½
            logger.info(f"ğŸ“¦ å‘ç° BM25 æ¨¡å‹ç¼“å­˜: {cache_path}ï¼Œä½¿ç”¨ç¦»çº¿æ¨¡å¼åŠ è½½")
            self._encoder = SparseTextEmbedding(model_name=self.model_name)
            logger.info(f"âœ… åˆå§‹åŒ– BM25 ç¨€ç–ç¼–ç å™¨: {self.model_name}")
        else:
            # æ— æœ¬åœ°ç¼“å­˜æ—¶ï¼Œä¸´æ—¶ç¦ç”¨ç¦»çº¿æ¨¡å¼ä»¥ä¸‹è½½æ¨¡å‹
            offline_vars = ['HF_HUB_OFFLINE', 'TRANSFORMERS_OFFLINE', 'HF_DATASETS_OFFLINE']
            old_values = {var: os.environ.pop(var, None) for var in offline_vars}
            
            try:
                import huggingface_hub.constants
                if hasattr(huggingface_hub.constants, 'HF_HUB_OFFLINE'):
                    huggingface_hub.constants.HF_HUB_OFFLINE = False
            except Exception:
                pass
            
            try:
                self._encoder = SparseTextEmbedding(model_name=self.model_name)
                logger.info(f"âœ… åˆå§‹åŒ– BM25 ç¨€ç–ç¼–ç å™¨: {self.model_name}")
            finally:
                for var, val in old_values.items():
                    if val is not None:
                        os.environ[var] = val

    def encode_documents(self, texts: List[str]) -> List:
        """ç¼–ç æ–‡æ¡£åˆ—è¡¨"""
        return list(self._encoder.embed(texts))

    def encode_query(self, text: str):
        """ç¼–ç æŸ¥è¯¢"""
        results = list(self._encoder.query_embed(text))
        return results[0] if results else None


class CustomAPIEmbeddings(Embeddings):
    """è‡ªå®šä¹‰HTTP APIåµŒå…¥æœåŠ¡"""
    
    def __init__(self, api_base_url: str, api_key: str = None, custom_headers: dict = None, model_name: str = 'text-embedding'):
        self.api_base_url = api_base_url.rstrip('/')
        self.api_key = api_key
        self.custom_headers = custom_headers or {}
        self.model_name = model_name
        
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """åµŒå…¥å¤šä¸ªæ–‡æ¡£"""
        return [self.embed_query(text) for text in texts]
    
    def embed_query(self, text: str) -> List[float]:
        """åµŒå…¥å•ä¸ªæŸ¥è¯¢"""
        headers = {
            'Content-Type': 'application/json',
            **self.custom_headers
        }
        
        if self.api_key:
            headers['Authorization'] = f'Bearer {self.api_key}'
        
        data = {
            'input': text,
            'model': self.model_name  # ä½¿ç”¨é…ç½®çš„æ¨¡å‹å
        }
        
        try:
            response = requests.post(
                self.api_base_url,  # ç›´æ¥ä½¿ç”¨å®Œæ•´çš„API URL
                json=data,
                headers=headers,
                timeout=30
            )
            response.raise_for_status()
            
            result = response.json()
            if 'data' in result and len(result['data']) > 0:
                return result['data'][0]['embedding']
            else:
                raise ValueError(f"APIå“åº”æ ¼å¼é”™è¯¯: {result}")
                
        except Exception as e:
            raise RuntimeError(f"è‡ªå®šä¹‰APIåµŒå…¥å¤±è´¥: {str(e)}")




class DocumentProcessor:
    """æ–‡æ¡£å¤„ç†å™¨ - æ”¯æŒç»“æ„åŒ–è§£æ"""

    def __init__(self):
        self.loaders = {
            'pdf': PyPDFLoader,
            'docx': self._load_docx_structured,  # ä½¿ç”¨è‡ªå®šä¹‰ç»“æ„åŒ–è§£æ
            'doc': self._load_doc_structured,    # æ”¯æŒæ—§ç‰ˆ .doc æ ¼å¼
            'pptx': UnstructuredPowerPointLoader,
            'txt': TextLoader,
            'md': UnstructuredMarkdownLoader,
            'html': UnstructuredHTMLLoader,
        }

    def load_document(self, document: Document) -> List[LangChainDocument]:
        """åŠ è½½æ–‡æ¡£å†…å®¹"""
        try:
            logger.info(f"å¼€å§‹åŠ è½½æ–‡æ¡£: {document.title} (ID: {document.id})")
            logger.info(f"æ–‡æ¡£ç±»å‹: {document.document_type}")

            # ä¼˜å…ˆçº§ï¼šURL > æ–‡æœ¬å†…å®¹ > æ–‡ä»¶
            if document.document_type == 'url' and document.url:
                logger.info(f"ä»URLåŠ è½½: {document.url}")
                return self._load_from_url(document.url)
            elif document.content:
                # å¦‚æœæœ‰æ–‡æœ¬å†…å®¹ï¼Œç›´æ¥ä½¿ç”¨
                logger.info("ä»æ–‡æœ¬å†…å®¹åŠ è½½")
                return self._load_from_content(document.content, document.title)
            elif document.file and hasattr(document.file, 'path'):
                file_path = document.file.path
                logger.info(f"ä»æ–‡ä»¶åŠ è½½: {file_path}")

                # Windowsè·¯å¾„å…¼å®¹æ€§å¤„ç†
                if os.name == 'nt':  # Windowsç³»ç»Ÿ
                    file_path = os.path.normpath(file_path)
                    if not os.path.isabs(file_path):
                        file_path = os.path.abspath(file_path)

                # æ£€æŸ¥æ–‡ä»¶æ˜¯å¦å­˜åœ¨
                if os.path.exists(file_path):
                    logger.info(f"æ–‡ä»¶å­˜åœ¨ï¼Œå¼€å§‹åŠ è½½: {file_path}")
                    return self._load_from_file(document)
                else:
                    raise FileNotFoundError(f"æ–‡ä»¶ä¸å­˜åœ¨: {file_path}")
            else:
                raise ValueError("æ–‡æ¡£æ²¡æœ‰å¯ç”¨çš„å†…å®¹æºï¼ˆæ— URLã€æ— æ–‡æœ¬å†…å®¹ã€æ— æ–‡ä»¶ï¼‰")
        except Exception as e:
            logger.error(f"åŠ è½½æ–‡æ¡£å¤±è´¥ {document.id}: {e}")
            raise

    def _load_from_url(self, url: str) -> List[LangChainDocument]:
        """ä»URLåŠ è½½æ–‡æ¡£"""
        loader = WebBaseLoader(url)
        return loader.load()

    def _load_from_content(self, content: str, title: str) -> List[LangChainDocument]:
        """ä»æ–‡æœ¬å†…å®¹åŠ è½½æ–‡æ¡£"""
        return [LangChainDocument(
            page_content=content,
            metadata={"source": title, "title": title}
        )]

    def _load_from_file(self, document: Document) -> List[LangChainDocument]:
        """ä»æ–‡ä»¶åŠ è½½æ–‡æ¡£"""
        file_path = document.file.path

        # Windowsè·¯å¾„å…¼å®¹æ€§å¤„ç†
        if os.name == 'nt':  # Windowsç³»ç»Ÿ
            # ç¡®ä¿è·¯å¾„ä½¿ç”¨æ­£ç¡®çš„åˆ†éš”ç¬¦
            file_path = os.path.normpath(file_path)
            # è½¬æ¢ä¸ºç»å¯¹è·¯å¾„
            if not os.path.isabs(file_path):
                file_path = os.path.abspath(file_path)

        logger.info(f"å°è¯•åŠ è½½æ–‡ä»¶: {file_path}")

        # å†æ¬¡æ£€æŸ¥æ–‡ä»¶æ˜¯å¦å­˜åœ¨
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"æ–‡ä»¶ä¸å­˜åœ¨: {file_path}")

        loader = self.loaders.get(document.document_type)
        if not loader:
            raise ValueError(f"ä¸æ”¯æŒçš„æ–‡æ¡£ç±»å‹: {document.document_type}")

        try:
            # æ£€æŸ¥æ˜¯å¦ä¸ºè‡ªå®šä¹‰æ–¹æ³•ï¼ˆdocx/doc ç»“æ„åŒ–è§£æï¼‰
            if callable(loader) and hasattr(loader, '__self__'):
                docs = loader(file_path, document)
            elif document.document_type == 'txt':
                # å¯¹äºæ–‡æœ¬æ–‡ä»¶ï¼Œä½¿ç”¨UTF-8ç¼–ç 
                loader_instance = loader(file_path, encoding='utf-8')
                docs = loader_instance.load()
            else:
                # å…¶ä»–ç±»å‹ä½¿ç”¨æ ‡å‡† LangChain loader
                loader_instance = loader(file_path)
                docs = loader_instance.load()

            # æ£€æŸ¥æ˜¯å¦æˆåŠŸåŠ è½½å†…å®¹
            if not docs:
                raise ValueError(f"æ–‡æ¡£åŠ è½½å¤±è´¥ï¼Œæ²¡æœ‰å†…å®¹: {file_path}")

            logger.info(f"æˆåŠŸåŠ è½½æ–‡æ¡£ï¼Œé¡µæ•°: {len(docs)}")

            # æ·»åŠ å…ƒæ•°æ®
            for doc in docs:
                doc.metadata.update({
                    "source": document.title,
                    "document_id": str(document.id),
                    "document_type": document.document_type,
                    "title": document.title,
                    "file_path": file_path
                })

            return docs

        except Exception as e:
            logger.error(f"æ–‡æ¡£åŠ è½½å™¨å¤±è´¥: {e}")
            # å¦‚æœæ˜¯æ–‡æœ¬æ–‡ä»¶ï¼Œå°è¯•ç›´æ¥è¯»å–
            if document.document_type == 'txt':
                try:
                    logger.info("å°è¯•ç›´æ¥è¯»å–æ–‡æœ¬æ–‡ä»¶...")
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()

                    if not content.strip():
                        raise ValueError("æ–‡ä»¶å†…å®¹ä¸ºç©º")

                    return [LangChainDocument(
                        page_content=content,
                        metadata={
                            "source": document.title,
                            "document_id": str(document.id),
                            "document_type": document.document_type,
                            "title": document.title,
                            "file_path": file_path
                        }
                    )]
                except Exception as read_error:
                    logger.error(f"ç›´æ¥è¯»å–æ–‡ä»¶ä¹Ÿå¤±è´¥: {read_error}")
                    raise
            else:
                raise

    def _load_docx_structured(self, file_path: str, document: Document) -> List[LangChainDocument]:
        """ç»“æ„åŒ–è§£æ .docx æ–‡ä»¶ï¼Œä¿ç•™æ ‡é¢˜å±‚çº§å’Œè¡¨æ ¼ç»“æ„"""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            logger.info(f"å¼€å§‹ç»“æ„åŒ–è§£æ Word æ–‡æ¡£ï¼Œæ®µè½æ•°: {len(doc.paragraphs)}, è¡¨æ ¼æ•°: {len(doc.tables)}")

            # åˆ›å»ºå…ƒç´ åˆ°å¯¹è±¡çš„æ˜ å°„
            paragraph_map = {p._element: p for p in doc.paragraphs}
            table_map = {t._element: t for t in doc.tables}

            content_parts = []
            extracted_paragraphs = 0
            extracted_tables = 0

            # æŒ‰æ–‡æ¡£é¡ºåºéå†æ‰€æœ‰å…ƒç´ 
            for element in doc.element.body:
                if element.tag.endswith('p'):  # æ®µè½
                    paragraph = paragraph_map.get(element)
                    if paragraph:
                        text = paragraph.text.strip()
                        if text:
                            markdown_text = self._convert_paragraph_to_markdown(paragraph)
                            content_parts.append(markdown_text)
                            extracted_paragraphs += 1

                elif element.tag.endswith('tbl'):  # è¡¨æ ¼
                    table = table_map.get(element)
                    if table:
                        table_content = self._extract_table_content(table)
                        if table_content:
                            content_parts.append(table_content)
                            extracted_tables += 1

            content = '\n\n'.join(content_parts)
            logger.info(f"Word ç»“æ„åŒ–è§£æå®Œæˆ - æ®µè½: {extracted_paragraphs}, è¡¨æ ¼: {extracted_tables}, å†…å®¹é•¿åº¦: {len(content)}")

            return [LangChainDocument(
                page_content=content,
                metadata={
                    "source": document.title,
                    "document_id": str(document.id),
                    "document_type": document.document_type,
                    "title": document.title,
                    "file_path": file_path,
                    "structured_parsing": True,
                    "paragraph_count": extracted_paragraphs,
                    "table_count": extracted_tables,
                }
            )]

        except Exception as e:
            logger.warning(f"ç»“æ„åŒ–è§£æå¤±è´¥ï¼Œé™çº§ä¸ºçº¯æ–‡æœ¬è§£æ: {e}")
            # é™çº§ä¸º Docx2txtLoader
            loader = Docx2txtLoader(file_path)
            docs = loader.load()
            for doc in docs:
                doc.metadata.update({
                    "source": document.title,
                    "document_id": str(document.id),
                    "document_type": document.document_type,
                    "title": document.title,
                    "file_path": file_path,
                    "structured_parsing": False,
                })
            return docs

    def _load_doc_structured(self, file_path: str, document: Document) -> List[LangChainDocument]:
        """è§£ææ—§ç‰ˆ .doc æ–‡ä»¶ï¼Œä¼˜å…ˆè½¬æ¢ä¸º docx ä»¥ä¿ç•™ç»“æ„"""
        import tempfile
        import subprocess

        # æ£€æµ‹æ–‡ä»¶çœŸå®æ ¼å¼
        with open(file_path, 'rb') as f:
            header = f.read(8)

        # ZIP é­”æ•°è¡¨ç¤ºå®é™…æ˜¯ .docx
        if header[:4] == b'PK\x03\x04':
            logger.info("æ£€æµ‹åˆ° .doc æ–‡ä»¶å®é™…ä¸º .docx æ ¼å¼ï¼Œä½¿ç”¨ docx è§£æå™¨")
            return self._load_docx_structured(file_path, document)

        try:
            # æ–¹æ³•1: LibreOffice è½¬æ¢ä¸º docx
            with tempfile.TemporaryDirectory() as tmp_dir:
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'docx',
                     '--outdir', tmp_dir, file_path],
                    capture_output=True, text=True, timeout=120
                )
                if result.returncode == 0:
                    import os
                    docx_file = os.path.join(tmp_dir,
                        os.path.basename(file_path).rsplit('.', 1)[0] + '.docx')
                    if os.path.exists(docx_file):
                        docs = self._load_docx_structured(docx_file, document)
                        logger.info(f"æˆåŠŸé€šè¿‡ LibreOffice è½¬æ¢å¹¶è§£æ .doc æ–‡ä»¶")
                        return docs
        except FileNotFoundError:
            logger.debug("LibreOffice æœªå®‰è£…")
        except Exception as e:
            logger.debug(f"LibreOffice è½¬æ¢å¤±è´¥: {e}")

        # æ–¹æ³•2: antiword æå–çº¯æ–‡æœ¬å¹¶æ¨æ–­æ ‡é¢˜
        try:
            result = subprocess.run(
                ['antiword', '-w', '0', file_path],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                content = self._infer_headings_from_plain_text(result.stdout.strip())
                logger.info(f"æˆåŠŸé€šè¿‡ antiword è§£æ .doc æ–‡ä»¶ï¼Œå†…å®¹é•¿åº¦: {len(content)}")
                return [LangChainDocument(
                    page_content=content,
                    metadata={
                        "source": document.title,
                        "document_id": str(document.id),
                        "document_type": document.document_type,
                        "title": document.title,
                        "file_path": file_path,
                        "structured_parsing": False,
                        "parse_method": "antiword",
                    }
                )]
        except FileNotFoundError:
            logger.debug("antiword æœªå®‰è£…")
        except Exception as e:
            logger.debug(f"antiword å¤±è´¥: {e}")

        raise ValueError(
            "æ— æ³•è§£æ .doc æ–‡ä»¶ã€‚è¯·å®‰è£… LibreOffice ä»¥è·å¾—æœ€ä½³æ•ˆæœï¼š\n"
            "Ubuntu/Debian: apt-get install libreoffice\n"
            "æˆ–è€…å°†æ–‡ä»¶å¦å­˜ä¸º .docx æ ¼å¼åé‡æ–°ä¸Šä¼ "
        )

    def _convert_paragraph_to_markdown(self, paragraph) -> str:
        """å°† Word æ®µè½è½¬æ¢ä¸º Markdown æ ¼å¼"""
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = paragraph.style.name if paragraph.style else ""

        # æ ¹æ®æ ·å¼è½¬æ¢ä¸º Markdown æ ‡é¢˜
        heading_map = {
            'heading 1': '# ', 'heading 2': '## ', 'heading 3': '### ',
            'heading 4': '#### ', 'heading 5': '##### ', 'heading 6': '###### ',
        }

        for style_key, prefix in heading_map.items():
            if style_key in style_name.lower():
                return f"{prefix}{text}"

        return text

    def _extract_table_content(self, table, depth=0) -> str:
        """æå–è¡¨æ ¼å†…å®¹ä¸º Markdown æ ¼å¼ï¼Œæ”¯æŒåˆå¹¶å•å…ƒæ ¼"""
        try:
            # è·å–è¡¨æ ¼çš„å®é™…è¡Œåˆ—æ•°
            row_count = len(table.rows)
            if row_count == 0:
                return ""

            # ä¼˜å…ˆç”¨ table.columns è·å–åˆ—æ•°ï¼ˆæ›´ç¨³å®šï¼‰ï¼Œå¤±è´¥æ—¶å›é€€åˆ°è¡Œ cells çš„æœ€å¤§é•¿åº¦
            try:
                max_cols = len(table.columns)
            except Exception:
                max_cols = max((len(row.cells) for row in table.rows), default=0)
            if max_cols == 0:
                return ""

            # æ„å»ºå•å…ƒæ ¼ç½‘æ ¼ï¼Œç”¨äºå¤„ç†åˆå¹¶å•å…ƒæ ¼
            # grid[row][col] = (cell_text, is_merged_continuation)
            grid = [[("", False) for _ in range(max_cols)] for _ in range(row_count)]

            def _sanitize_cell_text(text: str) -> str:
                """æ¸…ç†å•å…ƒæ ¼æ–‡æœ¬ï¼Œè½¬ä¹‰ Markdown ç‰¹æ®Šå­—ç¬¦"""
                text = (text or "").replace("\r", " ").replace("\n", " ").replace("\t", " ")
                text = " ".join(text.split())  # åˆå¹¶å¤šä¸ªç©ºæ ¼
                return text.replace("|", "\\|")  # è½¬ä¹‰ç®¡é“ç¬¦

            for row_idx, row in enumerate(table.rows):
                # è®°å½•å½“å‰è¡Œå·²å¤„ç†çš„å•å…ƒæ ¼ï¼ˆç”¨äºæ£€æµ‹æ°´å¹³åˆå¹¶ï¼šåŒä¸€è¡Œä¸­é‡å¤å¼•ç”¨åŒä¸€ tcï¼‰
                processed_cells_in_row = set()

                for col_idx, cell in enumerate(row.cells):
                    if col_idx >= max_cols:
                        break

                    # è·å–å•å…ƒæ ¼çš„å”¯ä¸€æ ‡è¯†
                    cell_id = id(cell._tc)

                    # æ£€æµ‹å‚ç›´åˆå¹¶
                    # Word ä¸­ <w:vMerge/> æ—  val å±æ€§æ—¶é€šå¸¸è¡¨ç¤º"ç»§ç»­åˆå¹¶"
                    try:
                        tcPr = cell._tc.tcPr
                        v_merge = tcPr.vMerge if tcPr is not None else None
                        v_merge_val = getattr(v_merge, "val", None) if v_merge is not None else None

                        if v_merge is None:
                            is_v_merge_continue = False
                        elif v_merge_val == "restart":
                            is_v_merge_continue = False
                        elif v_merge_val == "continue":
                            is_v_merge_continue = True
                        else:
                            # val ä¸º None æ—¶ï¼Œé€šè¿‡æ£€æŸ¥ä¸Šä¸€è¡Œæ˜¯å¦æœ‰ vMerge æ¥åˆ¤æ–­
                            if row_idx == 0:
                                is_v_merge_continue = False
                            else:
                                try:
                                    prev_cell = table.rows[row_idx - 1].cells[col_idx]
                                    prev_tcPr = prev_cell._tc.tcPr
                                    prev_v_merge = prev_tcPr.vMerge if prev_tcPr is not None else None
                                    is_v_merge_continue = prev_v_merge is not None
                                except Exception:
                                    is_v_merge_continue = False
                    except Exception:
                        is_v_merge_continue = False

                    # å‚ç›´åˆå¹¶ç»§ç»­ï¼šæ ‡è®°ä¸ºå·²åˆå¹¶
                    if is_v_merge_continue:
                        grid[row_idx][col_idx] = ("", True)
                        continue

                    # æ°´å¹³åˆå¹¶ç»§ç»­ï¼šåŒä¸€è¡Œä¸­é‡å¤å¼•ç”¨åŒä¸€ tc
                    if cell_id in processed_cells_in_row:
                        grid[row_idx][col_idx] = ("", True)
                        continue

                    # æå–å•å…ƒæ ¼å†…å®¹
                    nested_tables = cell.tables
                    if nested_tables and depth < 3:
                        cell_text_parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                        for nested_table in nested_tables:
                            nested_content = self._extract_table_content(nested_table, depth + 1)
                            if nested_content:
                                cell_text_parts.append(f"[åµŒå¥—è¡¨æ ¼] {nested_content}")
                        cell_text = _sanitize_cell_text(" ".join(cell_text_parts))
                    else:
                        cell_text = _sanitize_cell_text(cell.text.strip())

                    # å¡«å……ç½‘æ ¼
                    grid[row_idx][col_idx] = (cell_text, False)
                    processed_cells_in_row.add(cell_id)

            # ç”Ÿæˆ Markdown è¡¨æ ¼
            table_rows = []
            for row_idx in range(row_count):
                row_cells = [grid[row_idx][col_idx][0] for col_idx in range(max_cols)]

                # è·³è¿‡å…¨ç©ºè¡Œ
                if not any(cell.strip() for cell in row_cells):
                    continue

                table_rows.append(" | ".join(row_cells))

                # ç¬¬ä¸€è¡Œåæ·»åŠ åˆ†éš”ç¬¦
                if len(table_rows) == 1:
                    separator = " | ".join(["---"] * max_cols)
                    table_rows.append(separator)

            if table_rows:
                return "\n".join(table_rows)
            return ""

        except Exception as e:
            logger.warning(f"è¡¨æ ¼æå–å¤±è´¥: {e}")
            return ""

    def _infer_headings_from_plain_text(self, content: str) -> str:
        """ä»çº¯æ–‡æœ¬æ¨æ–­æ ‡é¢˜ç»“æ„"""
        import re

        lines = content.split('\n')
        result_lines = []

        # å¸¸è§æ ‡é¢˜æ¨¡å¼
        patterns = [
            (r'^ç¬¬[ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹åç™¾]+[ç« èŠ‚éƒ¨åˆ†]\s*', 1),
            (r'^[ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹å]+[ã€.ï¼]\s*', 2),
            (r'^[ï¼ˆ\(][ä¸€äºŒä¸‰å››äº”å…­ä¸ƒå…«ä¹å]+[ï¼‰\)]\s*', 3),
            (r'^(\d+)\s*[ã€.ï¼]\s*', 2),
            (r'^(\d+\.\d+)\s+', 3),
            (r'^(\d+\.\d+\.\d+)\s+', 4),
        ]

        for line in lines:
            stripped = line.strip()
            if not stripped:
                result_lines.append('')
                continue

            matched = False
            for pattern, level in patterns:
                if re.match(pattern, stripped) and len(stripped) <= 80:
                    prefix = '#' * level + ' '
                    result_lines.append(prefix + stripped)
                    matched = True
                    break

            if not matched:
                result_lines.append(line)

        return '\n'.join(result_lines)


class VectorStoreManager:
    """å‘é‡å­˜å‚¨ç®¡ç†å™¨ - æ”¯æŒç¨ å¯†+ç¨€ç–æ··åˆæ£€ç´¢"""

    # å‘é‡åç§°å¸¸é‡
    DENSE_VECTOR_NAME = "dense"
    SPARSE_VECTOR_NAME = "bm25"
    # RRF èåˆå‚æ•°
    RRF_K = 60

    # ç±»çº§åˆ«çš„ç¼“å­˜
    _vector_store_cache = {}
    _embeddings_cache = {}
    _sparse_encoder_cache = {}
    _global_config_cache = None
    _global_config_cache_time = 0

    def __init__(self, knowledge_base: KnowledgeBase):
        self.knowledge_base = knowledge_base
        self.global_config = self._get_global_config()
        self.embeddings = self._get_embeddings_instance()
        self.sparse_encoder = self._get_sparse_encoder()
        self._log_embedding_info()

    @classmethod
    def _get_global_config(cls):
        """è·å–å…¨å±€é…ç½®ï¼ˆå¸¦ç¼“å­˜ï¼Œ5åˆ†é’Ÿè¿‡æœŸï¼‰"""
        import time
        current_time = time.time()
        
        # ç¼“å­˜5åˆ†é’Ÿ
        if cls._global_config_cache and (current_time - cls._global_config_cache_time) < 300:
            return cls._global_config_cache
        
        cls._global_config_cache = KnowledgeGlobalConfig.get_config()
        cls._global_config_cache_time = current_time
        return cls._global_config_cache

    @classmethod
    def clear_global_config_cache(cls):
        """æ¸…ç†å…¨å±€é…ç½®ç¼“å­˜"""
        cls._global_config_cache = None
        cls._global_config_cache_time = 0

    def _get_embeddings_instance(self):
        """è·å–åµŒå…¥æ¨¡å‹å®ä¾‹ï¼Œä½¿ç”¨å…¨å±€é…ç½®"""
        config = self.global_config
        cache_key = f"{config.embedding_service}_{config.api_base_url}_{config.model_name}"
        
        if cache_key not in self._embeddings_cache:
            embedding_service = config.embedding_service
            
            try:
                if embedding_service == 'openai':
                    self._embeddings_cache[cache_key] = self._create_openai_embeddings(config)
                elif embedding_service == 'azure_openai':
                    self._embeddings_cache[cache_key] = self._create_azure_embeddings(config)
                elif embedding_service == 'ollama':
                    self._embeddings_cache[cache_key] = self._create_ollama_embeddings(config)
                elif embedding_service == 'custom':
                    self._embeddings_cache[cache_key] = self._create_custom_api_embeddings(config)
                else:
                    raise ValueError(f"ä¸æ”¯æŒçš„åµŒå…¥æœåŠ¡: {embedding_service}")
                    
                # æµ‹è¯•åµŒå…¥åŠŸèƒ½
                test_embedding = self._embeddings_cache[cache_key].embed_query("æ¨¡å‹åŠŸèƒ½æµ‹è¯•")
                logger.info(f"âœ… åµŒå…¥æ¨¡å‹æµ‹è¯•æˆåŠŸ: {embedding_service}, ç»´åº¦: {len(test_embedding)}")
                
            except Exception as e:
                logger.error(f"âŒ åµŒå…¥æœåŠ¡ {embedding_service} åˆå§‹åŒ–å¤±è´¥: {str(e)}")
                raise
                
        return self._embeddings_cache[cache_key]

    def _get_sparse_encoder(self) -> Optional[SparseBM25Encoder]:
        """è·å– BM25 ç¨€ç–ç¼–ç å™¨ï¼ˆå¸¦ç¼“å­˜ï¼‰"""
        cache_key = self.SPARSE_VECTOR_NAME
        
        if cache_key not in self._sparse_encoder_cache:
            try:
                self._sparse_encoder_cache[cache_key] = SparseBM25Encoder()
            except ImportError as e:
                logger.warning(f"âš ï¸ FastEmbed æœªå®‰è£…ï¼Œå°†ä½¿ç”¨çº¯ç¨ å¯†å‘é‡æ£€ç´¢: {e}")
                self._sparse_encoder_cache[cache_key] = None
            except Exception as e:
                logger.warning(f"âš ï¸ BM25 ç¼–ç å™¨åˆå§‹åŒ–å¤±è´¥: {e}ï¼Œé™çº§ä¸ºçº¯ç¨ å¯†æ£€ç´¢")
                self._sparse_encoder_cache[cache_key] = None
        
        return self._sparse_encoder_cache[cache_key]
    
    def _create_openai_embeddings(self, config):
        """åˆ›å»ºOpenAI Embeddingså®ä¾‹"""
        try:
            from langchain_openai import OpenAIEmbeddings
        except ImportError:
            raise ImportError("éœ€è¦å®‰è£…langchain-openai: pip install langchain-openai")
        
        kwargs = {
            'model': config.model_name or 'text-embedding-ada-002',
        }
        
        if config.api_key:
            kwargs['api_key'] = config.api_key
        if config.api_base_url:
            kwargs['base_url'] = config.api_base_url
            
        logger.info(f"ğŸš€ åˆå§‹åŒ–OpenAIåµŒå…¥æ¨¡å‹: {kwargs['model']}")
        return OpenAIEmbeddings(**kwargs)
    
    def _create_azure_embeddings(self, config):
        """åˆ›å»ºAzure OpenAI Embeddingså®ä¾‹"""
        try:
            from langchain_openai import AzureOpenAIEmbeddings
        except ImportError:
            raise ImportError("éœ€è¦å®‰è£…langchain-openai: pip install langchain-openai")
        
        if not all([config.api_key, config.api_base_url]):
            raise ValueError("Azure OpenAIéœ€è¦é…ç½®api_keyå’Œapi_base_url")
        
        kwargs = {
            'model': config.model_name or 'text-embedding-ada-002',
            'api_key': config.api_key,
            'azure_endpoint': config.api_base_url,
            'api_version': '2024-02-15-preview',
        }
        
        kwargs['deployment'] = config.model_name or 'text-embedding-ada-002'
            
        logger.info(f"ğŸš€ åˆå§‹åŒ–Azure OpenAIåµŒå…¥æ¨¡å‹: {kwargs['model']}")
        return AzureOpenAIEmbeddings(**kwargs)
    
    def _create_ollama_embeddings(self, config):
        """åˆ›å»ºOllama Embeddingså®ä¾‹"""
        try:
            from langchain_ollama import OllamaEmbeddings
        except ImportError:
            raise ImportError("éœ€è¦å®‰è£…langchain-ollama: pip install langchain-ollama")
        
        kwargs = {
            'model': config.model_name or 'nomic-embed-text',
        }
        
        if config.api_base_url:
            kwargs['base_url'] = config.api_base_url
        else:
            kwargs['base_url'] = 'http://localhost:11434'
            
        logger.info(f"ğŸš€ åˆå§‹åŒ–OllamaåµŒå…¥æ¨¡å‹: {kwargs['model']}")
        return OllamaEmbeddings(**kwargs)
    
    def _create_custom_api_embeddings(self, config):
        """åˆ›å»ºè‡ªå®šä¹‰API Embeddingså®ä¾‹"""
        if not config.api_base_url:
            raise ValueError("è‡ªå®šä¹‰APIéœ€è¦é…ç½®api_base_url")
        
        logger.info(f"ğŸš€ åˆå§‹åŒ–è‡ªå®šä¹‰APIåµŒå…¥æ¨¡å‹: {config.api_base_url}")
        return CustomAPIEmbeddings(
            api_base_url=config.api_base_url,
            api_key=config.api_key,
            custom_headers={},
            model_name=config.model_name
        )
    
    def _log_embedding_info(self):
        """è®°å½•åµŒå…¥æ¨¡å‹ä¿¡æ¯"""
        embedding_type = type(self.embeddings).__name__
        config = self.global_config
        logger.info(f"   ğŸŒŸ çŸ¥è¯†åº“: {self.knowledge_base.name}")
        logger.info(f"   ğŸ¯ é…ç½®çš„åµŒå…¥æ¨¡å‹: {config.model_name}")
        logger.info(f"   âœ… å®é™…ä½¿ç”¨çš„åµŒå…¥æ¨¡å‹: {embedding_type}")

        if embedding_type == "OpenAIEmbeddings":
            logger.info(f"   ğŸ‰ è¯´æ˜: ä½¿ç”¨OpenAIåµŒå…¥APIæœåŠ¡")
        elif embedding_type == "AzureOpenAIEmbeddings":
            logger.info(f"   ğŸ‰ è¯´æ˜: ä½¿ç”¨Azure OpenAIåµŒå…¥APIæœåŠ¡")
        elif embedding_type == "OllamaEmbeddings":
            logger.info(f"   ğŸ‰ è¯´æ˜: ä½¿ç”¨Ollamaæœ¬åœ°APIåµŒå…¥æœåŠ¡")
        elif embedding_type == "CustomAPIEmbeddings":
            logger.info(f"   ğŸ‰ è¯´æ˜: ä½¿ç”¨è‡ªå®šä¹‰HTTP APIåµŒå…¥æœåŠ¡")

        self._vector_store = None
        self._qdrant_client = None
        logger.info(f"ğŸ¤– å‘é‡å­˜å‚¨ç®¡ç†å™¨åˆå§‹åŒ–å®Œæˆ:")
        logger.info(f"   ğŸ“‹ çŸ¥è¯†åº“: {self.knowledge_base.name} (ID: {self.knowledge_base.id})")
        logger.info(f"   ğŸ¯ é…ç½®çš„åµŒå…¥æ¨¡å‹: {config.model_name}")
        logger.info(f"   âœ… å®é™…ä½¿ç”¨çš„åµŒå…¥æ¨¡å‹: {embedding_type}")
        logger.info(f"   ğŸ’¾ å‘é‡å­˜å‚¨ç±»å‹: Qdrant")

    def _get_qdrant_url(self) -> str:
        """è·å– Qdrant æœåŠ¡åœ°å€"""
        return os.environ.get('QDRANT_URL', 'http://localhost:8918')

    def _get_collection_name(self) -> str:
        """è·å–é›†åˆåç§°"""
        return f"kb_{self.knowledge_base.id}"

    @property
    def qdrant_client(self) -> QdrantClient:
        """è·å– Qdrant å®¢æˆ·ç«¯"""
        if self._qdrant_client is None:
            qdrant_url = self._get_qdrant_url()
            self._qdrant_client = QdrantClient(url=qdrant_url)
            logger.info(f"ğŸ”— å·²è¿æ¥ Qdrant: {qdrant_url}")
        return self._qdrant_client

    @property
    def vector_store(self):
        """è·å–å‘é‡å­˜å‚¨å®ä¾‹ï¼ˆå¸¦ç¼“å­˜å’Œå¥åº·æ£€æŸ¥ï¼‰"""
        if self._vector_store is None:
            cache_key = str(self.knowledge_base.id)

            if cache_key in self._vector_store_cache:
                cached_store = self._vector_store_cache[cache_key]
                try:
                    # éªŒè¯ Qdrant é›†åˆæ˜¯å¦å­˜åœ¨
                    self.qdrant_client.get_collection(self._get_collection_name())
                    logger.info(f"ä½¿ç”¨ç¼“å­˜çš„å‘é‡å­˜å‚¨å®ä¾‹: {cache_key}")
                    self._vector_store = cached_store
                except Exception as e:
                    logger.warning(f"ç¼“å­˜çš„ Collection æ— æ•ˆ,é‡æ–°åˆ›å»º: {e}")
                    del self._vector_store_cache[cache_key]
                    logger.info(f"åˆ›å»ºæ–°çš„å‘é‡å­˜å‚¨å®ä¾‹: {cache_key}")
                    self._vector_store = self._create_vector_store()
                    self._vector_store_cache[cache_key] = self._vector_store
            else:
                logger.info(f"åˆ›å»ºæ–°çš„å‘é‡å­˜å‚¨å®ä¾‹: {cache_key}")
                self._vector_store = self._create_vector_store()
                self._vector_store_cache[cache_key] = self._vector_store

        return self._vector_store

    @classmethod
    def clear_cache(cls, knowledge_base_id=None):
        """æ¸…ç†å‘é‡å­˜å‚¨ç¼“å­˜"""
        if knowledge_base_id:
            cache_key = str(knowledge_base_id)
            if cache_key in cls._vector_store_cache:
                del cls._vector_store_cache[cache_key]
                logger.info(f"å·²æ¸…ç†çŸ¥è¯†åº“ {cache_key} çš„å‘é‡å­˜å‚¨ç¼“å­˜")

            # æ¸…ç† Qdrant é›†åˆ
            try:
                qdrant_url = os.environ.get('QDRANT_URL', 'http://localhost:8918')
                client = QdrantClient(url=qdrant_url)
                collection_name = f"kb_{knowledge_base_id}"
                if client.collection_exists(collection_name):
                    client.delete_collection(collection_name)
                    logger.info(f"å·²åˆ é™¤ Qdrant é›†åˆ: {collection_name}")
            except Exception as e:
                logger.warning(f"æ¸…ç† Qdrant é›†åˆå¤±è´¥: {e}")
        else:
            # æ¸…ç†æ‰€æœ‰ç¼“å­˜
            cls._vector_store_cache.clear()
            cls._embeddings_cache.clear()
            cls._sparse_encoder_cache.clear()
            logger.info("å·²æ¸…ç†æ‰€æœ‰å‘é‡å­˜å‚¨ç¼“å­˜")

    def _create_vector_store(self):
        """åˆ›å»º Qdrant å‘é‡å­˜å‚¨ï¼ˆæ”¯æŒç¨ å¯†+ç¨€ç–æ··åˆï¼‰"""
        collection_name = self._get_collection_name()
        
        # è·å–åµŒå…¥å‘é‡ç»´åº¦
        test_embedding = self.embeddings.embed_query("æµ‹è¯•")
        vector_size = len(test_embedding)
        
        # é…ç½®å‘½åå‘é‡ï¼ˆç”¨äºæ··åˆæ£€ç´¢ï¼‰
        vectors_config = {
            self.DENSE_VECTOR_NAME: VectorParams(
                size=vector_size,
                distance=Distance.COSINE
            )
        }
        
        # é…ç½®ç¨€ç–å‘é‡
        sparse_vectors_config = None
        if self.sparse_encoder:
            sparse_vectors_config = {
                self.SPARSE_VECTOR_NAME: SparseVectorParams(
                    index=SparseIndexParams(on_disk=False)
                )
            }
        
        # ç¡®ä¿é›†åˆå­˜åœ¨
        try:
            if not self.qdrant_client.collection_exists(collection_name):
                self.qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=vectors_config,
                    sparse_vectors_config=sparse_vectors_config
                )
                mode = "ç¨€ç–+ç¨ å¯†æ··åˆ" if sparse_vectors_config else "çº¯ç¨ å¯†"
                logger.info(f"âœ… åˆ›å»º Qdrant é›†åˆ: {collection_name}, ç»´åº¦: {vector_size}, æ¨¡å¼: {mode}")
            else:
                # æ£€æŸ¥æ˜¯å¦éœ€è¦æ›´æ–°ç¨€ç–é…ç½®
                if sparse_vectors_config:
                    try:
                        self.qdrant_client.update_collection(
                            collection_name=collection_name,
                            sparse_vectors_config=sparse_vectors_config
                        )
                    except Exception as e:
                        logger.debug(f"è·³è¿‡ç¨€ç–é…ç½®æ›´æ–°: {e}")
        except Exception as e:
            logger.warning(f"æ£€æŸ¥/åˆ›å»ºé›†åˆæ—¶å‡ºé”™: {e}")
        
        # ä½¿ç”¨ LangChain çš„ QdrantVectorStoreï¼ˆç”¨äºå…¼å®¹æ€§ï¼Œå®é™…æ··åˆæŸ¥è¯¢ç›´æ¥ç”¨ clientï¼‰
        qdrant_store = QdrantVectorStore(
            client=self.qdrant_client,
            collection_name=collection_name,
            embedding=self.embeddings,
            vector_name=self.DENSE_VECTOR_NAME
        )
        
        return qdrant_store

    def add_documents(self, documents: List[LangChainDocument], document_obj: Document) -> List[str]:
        """æ·»åŠ æ–‡æ¡£åˆ°å‘é‡å­˜å‚¨ï¼ˆç¨ å¯†+ç¨€ç–æ··åˆï¼‰"""
        try:
            # ç¡®ä¿é›†åˆå­˜åœ¨ï¼ˆè§¦å‘ vector_store å±æ€§ä¼šåˆ›å»ºé›†åˆï¼‰
            _ = self.vector_store
            
            # æ–‡æ¡£åˆ†å—
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.knowledge_base.chunk_size,
                chunk_overlap=self.knowledge_base.chunk_overlap
            )
            chunks = text_splitter.split_documents(documents)
            
            # ç”Ÿæˆå”¯ä¸€çš„ vector_ids
            vector_ids = [str(uuid.uuid4()) for _ in chunks]
            chunk_texts = [chunk.page_content for chunk in chunks]
            
            # è®¡ç®—ç¨ å¯†å‘é‡
            dense_embeddings = self.embeddings.embed_documents(chunk_texts)
            
            # è®¡ç®—ç¨€ç–å‘é‡ï¼ˆå¦‚æœå¯ç”¨ï¼‰
            sparse_embeddings = None
            if self.sparse_encoder:
                sparse_embeddings = self.sparse_encoder.encode_documents(chunk_texts)
            
            # æ„å»º PointStruct åˆ—è¡¨
            points: List[PointStruct] = []
            for i, (chunk, vector_id, dense_vector) in enumerate(zip(chunks, vector_ids, dense_embeddings)):
                payload = dict(chunk.metadata or {})
                payload.update({
                    "page_content": chunk.page_content,
                    "document_id": str(document_obj.id),
                    "chunk_index": i,
                    "vector_id": vector_id,
                    "knowledge_base_id": str(self.knowledge_base.id),
                })
                
                # æ„å»ºå‘é‡é…ç½®
                vectors = {self.DENSE_VECTOR_NAME: dense_vector}
                
                # æ·»åŠ ç¨€ç–å‘é‡ï¼ˆå¦‚æœå¯ç”¨ï¼‰
                sparse_vectors = None
                if sparse_embeddings and sparse_embeddings[i]:
                    sparse_vec = sparse_embeddings[i]
                    sparse_vectors = {
                        self.SPARSE_VECTOR_NAME: SparseVector(
                            indices=sparse_vec.indices.tolist(),
                            values=sparse_vec.values.tolist(),
                        )
                    }
                
                point = PointStruct(
                    id=vector_id,
                    vector=vectors,
                    payload=payload,
                )
                
                # Qdrant SDK éœ€è¦å•ç‹¬è®¾ç½® sparse_vectors
                if sparse_vectors:
                    point = PointStruct(
                        id=vector_id,
                        vector={
                            self.DENSE_VECTOR_NAME: dense_vector,
                            self.SPARSE_VECTOR_NAME: SparseVector(
                                indices=sparse_embeddings[i].indices.tolist(),
                                values=sparse_embeddings[i].values.tolist(),
                            )
                        },
                        payload=payload,
                    )
                
                points.append(point)
            
            # æ‰¹é‡å†™å…¥ Qdrant
            self.qdrant_client.upsert(
                collection_name=self._get_collection_name(),
                points=points,
            )
            
            mode = "ç¨€ç–+ç¨ å¯†" if sparse_embeddings else "çº¯ç¨ å¯†"
            logger.info(f"âœ… å·²å†™å…¥ {len(points)} ä¸ªåˆ†å—åˆ° Qdrantï¼ˆ{mode}ï¼‰")

            # ä¿å­˜åˆ†å—ä¿¡æ¯åˆ°æ•°æ®åº“
            self._save_chunks_to_db(chunks, vector_ids, document_obj)

            return vector_ids
        except Exception as e:
            logger.error(f"æ·»åŠ æ–‡æ¡£åˆ°å‘é‡å­˜å‚¨å¤±è´¥: {e}")
            raise

    def _save_chunks_to_db(self, chunks: List[LangChainDocument], vector_ids: List[str], document_obj: Document):
        """ä¿å­˜åˆ†å—ä¿¡æ¯åˆ°æ•°æ®åº“"""
        chunk_objects = []
        for i, (chunk, vector_id) in enumerate(zip(chunks, vector_ids)):
            # è®¡ç®—å†…å®¹å“ˆå¸Œ
            content_hash = hashlib.md5(chunk.page_content.encode()).hexdigest()

            chunk_obj = DocumentChunk(
                document=document_obj,
                chunk_index=i,
                content=chunk.page_content,
                vector_id=vector_id,
                embedding_hash=content_hash,
                start_index=chunk.metadata.get('start_index'),
                end_index=chunk.metadata.get('end_index'),
                page_number=chunk.metadata.get('page')
            )
            chunk_objects.append(chunk_obj)

        DocumentChunk.objects.bulk_create(chunk_objects)

    def similarity_search(self, query: str, k: int = 5, score_threshold: float = 0.1) -> List[Dict[str, Any]]:
        """ç›¸ä¼¼åº¦æœç´¢ï¼ˆæ”¯æŒç¨ å¯†+ç¨€ç–æ··åˆæ£€ç´¢ï¼‰"""
        embedding_type = type(self.embeddings).__name__
        logger.info(f"ğŸ” å¼€å§‹ç›¸ä¼¼åº¦æœç´¢ (Qdrant):")
        logger.info(f"   ğŸ“ æŸ¥è¯¢: '{query}'")
        logger.info(f"   ğŸ¤– ä½¿ç”¨åµŒå…¥æ¨¡å‹: {embedding_type}")
        logger.info(f"   ğŸ¯ è¿”å›æ•°é‡: {k}, ç›¸ä¼¼åº¦é˜ˆå€¼: {score_threshold}")

        # æ ¹æ®æ˜¯å¦æœ‰ç¨€ç–ç¼–ç å™¨é€‰æ‹©æ£€ç´¢æ–¹å¼
        if self.sparse_encoder:
            logger.info("   ğŸ”€ ä½¿ç”¨æ··åˆæ£€ç´¢ï¼ˆBM25 + ç¨ å¯†å‘é‡ï¼‰")
            return self._hybrid_similarity_search(query, k, score_threshold)
        else:
            logger.info("   ğŸ“Š ä½¿ç”¨çº¯ç¨ å¯†å‘é‡æ£€ç´¢")
            return self._dense_similarity_search(query, k, score_threshold)

    def _dense_similarity_search(self, query: str, k: int, score_threshold: float) -> List[Dict[str, Any]]:
        """çº¯ç¨ å¯†å‘é‡æ£€ç´¢"""
        try:
            dense_vector = self.embeddings.embed_query(query)
            collection_name = self._get_collection_name()
            
            results = self.qdrant_client.search(
                collection_name=collection_name,
                query_vector=NamedVector(
                    name=self.DENSE_VECTOR_NAME,
                    vector=dense_vector,
                ),
                limit=k,
                with_payload=True,
            )
            
            logger.info(f"ğŸ” ç¨ å¯†æ£€ç´¢ç»“æœ: {len(results)}")
            return self._format_search_results(results, score_threshold)
            
        except Exception as e:
            logger.error(f"ç¨ å¯†å‘é‡æœç´¢å¤±è´¥: {e}")
            raise

    def _hybrid_similarity_search(self, query: str, k: int, score_threshold: float) -> List[Dict[str, Any]]:
        """æ··åˆæ£€ç´¢ï¼ˆRRF èåˆç¨ å¯†+ç¨€ç–ï¼‰"""
        try:
            collection_name = self._get_collection_name()
            per_source_limit = max(k * 3, 10)  # æ¯ç§æ£€ç´¢æ–¹å¼å¤šå–ä¸€äº›å€™é€‰
            
            # è®¡ç®—ç¨ å¯†å‘é‡
            dense_vector = self.embeddings.embed_query(query)
            
            # è®¡ç®—ç¨€ç–å‘é‡
            sparse_query = self.sparse_encoder.encode_query(query)
            
            # ç¨ å¯†å‘é‡æ£€ç´¢
            dense_results = self.qdrant_client.search(
                collection_name=collection_name,
                query_vector=NamedVector(
                    name=self.DENSE_VECTOR_NAME,
                    vector=dense_vector,
                ),
                limit=per_source_limit,
                with_payload=True,
            )
            
            # ç¨€ç–å‘é‡æ£€ç´¢
            sparse_results = []
            if sparse_query:
                sparse_results = self.qdrant_client.search(
                    collection_name=collection_name,
                    query_vector=NamedSparseVector(
                        name=self.SPARSE_VECTOR_NAME,
                        vector=SparseVector(
                            indices=sparse_query.indices.tolist(),
                            values=sparse_query.values.tolist(),
                        ),
                    ),
                    limit=per_source_limit,
                    with_payload=True,
                )
            
            logger.info(f"ğŸ” ç¨ å¯†å€™é€‰: {len(dense_results)}, ç¨€ç–å€™é€‰: {len(sparse_results)}")
            
            # RRF èåˆ
            fused_results = self._rrf_fusion(dense_results, sparse_results, k)
            
            return self._format_fused_results(fused_results, score_threshold)
            
        except Exception as e:
            logger.error(f"æ··åˆæœç´¢å¤±è´¥: {e}")
            # é™çº§ä¸ºçº¯ç¨ å¯†æ£€ç´¢
            logger.warning("âš ï¸ é™çº§ä¸ºçº¯ç¨ å¯†æ£€ç´¢")
            return self._dense_similarity_search(query, k, score_threshold)

    def _rrf_fusion(self, dense_results, sparse_results, limit: int) -> List[Dict[str, Any]]:
        """RRF (Reciprocal Rank Fusion) èåˆä¸¤ç§æ£€ç´¢ç»“æœ"""
        if not dense_results and not sparse_results:
            return []
        
        fused: Dict[str, Dict[str, Any]] = {}
        contributors = 0
        
        def accumulate(results, label: str):
            for rank, point in enumerate(results):
                point_id = str(point.id)
                if point_id not in fused:
                    fused[point_id] = {
                        "payload": point.payload or {},
                        "score": 0.0,
                        "labels": {},
                        "original_scores": {},
                    }
                incremental = 1.0 / (self.RRF_K + rank + 1)
                fused[point_id]["score"] += incremental
                fused[point_id]["labels"][label] = incremental
                fused[point_id]["original_scores"][label] = point.score
        
        if dense_results:
            contributors += 1
            accumulate(dense_results, "dense")
        if sparse_results:
            contributors += 1
            accumulate(sparse_results, "sparse")
        
        # å½’ä¸€åŒ–åˆ†æ•°åˆ° 0-1 èŒƒå›´
        max_possible = contributors * (1.0 / (self.RRF_K + 1))
        max_possible = max(max_possible, 1e-9)
        
        fused_list = []
        for point_id, data in fused.items():
            data["id"] = point_id
            data["score"] = min(data["score"] / max_possible, 1.0)
            fused_list.append(data)
        
        # æŒ‰èåˆåˆ†æ•°é™åºæ’åº
        fused_list.sort(key=lambda item: item["score"], reverse=True)
        return fused_list[:limit]

    def _format_search_results(self, results, score_threshold: float) -> List[Dict[str, Any]]:
        """æ ¼å¼åŒ–ç¨ å¯†æœç´¢ç»“æœ"""
        formatted_results = []
        
        for i, point in enumerate(results):
            score = point.score
            if score < score_threshold:
                continue
            
            payload = point.payload or {}
            content = payload.get("page_content", "")
            
            result = {
                'content': content,
                'metadata': payload,
                'similarity_score': float(score)
            }
            formatted_results.append(result)
            
            source = payload.get('source', 'æœªçŸ¥æ¥æº')
            logger.info(f"   ğŸ“„ ç»“æœ{i+1}: ç›¸ä¼¼åº¦={score:.4f} ({score*100:.1f}%), æ¥æº={source}")
        
        # å¦‚æœæ²¡æœ‰æ»¡è¶³é˜ˆå€¼çš„ç»“æœï¼Œè¿”å›æœ€ä½³ç»“æœ
        if not formatted_results and results:
            best = results[0]
            payload = best.payload or {}
            formatted_results.append({
                'content': payload.get("page_content", ""),
                'metadata': payload,
                'similarity_score': float(best.score)
            })
        
        logger.info(f"ğŸ“Š è¿‡æ»¤åç»“æœæ•°é‡: {len(formatted_results)}")
        return formatted_results

    def _format_fused_results(self, fused_results: List[Dict], score_threshold: float) -> List[Dict[str, Any]]:
        """æ ¼å¼åŒ– RRF èåˆç»“æœ"""
        formatted_results = []
        
        for i, entry in enumerate(fused_results):
            score = entry["score"]
            if score < score_threshold:
                continue
            
            payload = entry.get("payload", {})
            content = payload.get("page_content", "")
            
            # æ·»åŠ èåˆæ¥æºä¿¡æ¯
            labels = entry.get("labels", {})
            original_scores = entry.get("original_scores", {})
            
            result = {
                'content': content,
                'metadata': payload,
                'similarity_score': float(score),
                'fusion_detail': {
                    'sources': list(labels.keys()),
                    'dense_score': original_scores.get("dense"),
                    'sparse_score': original_scores.get("sparse"),
                }
            }
            formatted_results.append(result)
            
            source = payload.get('source', 'æœªçŸ¥æ¥æº')
            sources_str = "+".join(labels.keys())
            logger.info(f"   ğŸ“„ ç»“æœ{i+1}: èåˆåˆ†={score:.4f} ({score*100:.1f}%), æ¥æº={source}, æ£€ç´¢æº=[{sources_str}]")
        
        # å¦‚æœæ²¡æœ‰æ»¡è¶³é˜ˆå€¼çš„ç»“æœï¼Œè¿”å›æœ€ä½³ç»“æœ
        if not formatted_results and fused_results:
            best = fused_results[0]
            payload = best.get("payload", {})
            formatted_results.append({
                'content': payload.get("page_content", ""),
                'metadata': payload,
                'similarity_score': float(best["score"]),
            })
        
        logger.info(f"ğŸ“Š è¿‡æ»¤åç»“æœæ•°é‡: {len(formatted_results)}")
        return formatted_results

    def delete_document(self, document: Document):
        """ä» Qdrant å‘é‡å­˜å‚¨ä¸­åˆ é™¤æ–‡æ¡£"""
        try:
            chunks = document.chunks.all()
            vector_ids = [chunk.vector_id for chunk in chunks if chunk.vector_id]

            if vector_ids:
                # Qdrant åˆ é™¤
                collection_name = self._get_collection_name()
                self.qdrant_client.delete(
                    collection_name=collection_name,
                    points_selector=vector_ids
                )
                logger.info(f"âœ… å·²ä» Qdrant åˆ é™¤ {len(vector_ids)} ä¸ªå‘é‡")

            chunks.delete()
        except Exception as e:
            logger.error(f"åˆ é™¤æ–‡æ¡£å‘é‡å¤±è´¥: {e}")
            raise


class KnowledgeBaseService:
    """çŸ¥è¯†åº“æœåŠ¡"""

    def __init__(self, knowledge_base: KnowledgeBase):
        self.knowledge_base = knowledge_base
        self.document_processor = DocumentProcessor()
        self.vector_manager = VectorStoreManager(knowledge_base)

    def process_document(self, document: Document) -> bool:
        """å¤„ç†æ–‡æ¡£"""
        try:
            # æ›´æ–°çŠ¶æ€ä¸ºå¤„ç†ä¸­
            document.status = 'processing'
            document.save()

            # æ¸…ç†å·²å­˜åœ¨çš„åˆ†å—å’Œå‘é‡ï¼ˆå¦‚æœæœ‰çš„è¯ï¼‰
            try:
                self.vector_manager.delete_document(document)
            except Exception as e:
                logger.warning(f"åˆ é™¤æ—§å‘é‡æ—¶å‡ºé”™ï¼ˆå¯èƒ½æ˜¯é¦–æ¬¡å¤„ç†ï¼‰: {e}")
            
            # å†ä»æ•°æ®åº“åˆ é™¤åˆ†å—è®°å½•
            document.chunks.all().delete()

            # åŠ è½½æ–‡æ¡£
            langchain_docs = self.document_processor.load_document(document)

            # è®¡ç®—æ–‡æ¡£ç»Ÿè®¡ä¿¡æ¯
            total_content = '\n'.join([doc.page_content for doc in langchain_docs])
            document.word_count = len(total_content.split())
            document.page_count = len(langchain_docs)

            # å‘é‡åŒ–å¹¶å­˜å‚¨
            vector_ids = self.vector_manager.add_documents(langchain_docs, document)

            # æ›´æ–°çŠ¶æ€ä¸ºå®Œæˆ
            document.status = 'completed'
            document.processed_at = timezone.now()
            document.error_message = None
            document.save()

            logger.info(f"æ–‡æ¡£å¤„ç†æˆåŠŸ: {document.id}, ç”Ÿæˆ {len(vector_ids)} ä¸ªåˆ†å—")
            return True

        except Exception as e:
            # æ›´æ–°çŠ¶æ€ä¸ºå¤±è´¥
            document.status = 'failed'
            document.error_message = str(e)
            document.save()

            logger.error(f"æ–‡æ¡£å¤„ç†å¤±è´¥: {document.id}, é”™è¯¯: {e}")
            return False

    def query(self, query_text: str, top_k: int = 5, similarity_threshold: float = 0.5,
              user=None) -> Dict[str, Any]:
        """æŸ¥è¯¢çŸ¥è¯†åº“"""
        start_time = time.time()

        try:
            # è®°å½•æŸ¥è¯¢å¼€å§‹ä¿¡æ¯
            embedding_type = type(self.vector_manager.embeddings).__name__
            logger.info(f"ğŸš€ çŸ¥è¯†åº“æŸ¥è¯¢å¼€å§‹:")
            logger.info(f"   ğŸ“š çŸ¥è¯†åº“: {self.knowledge_base.name}")
            logger.info(f"   ğŸ‘¤ ç”¨æˆ·: {user.username if user else 'åŒ¿å'}")
            logger.info(f"   ğŸ¤– åµŒå…¥æ¨¡å‹: {embedding_type}")
            logger.info(f"   ğŸ’¾ å‘é‡å­˜å‚¨: Qdrant")

            # æ‰§è¡Œæ£€ç´¢
            retrieval_start = time.time()
            search_results = self.vector_manager.similarity_search(
                query_text, k=top_k, score_threshold=similarity_threshold
            )
            retrieval_time = time.time() - retrieval_start

            # ç”Ÿæˆå›ç­”ï¼ˆè¿™é‡Œå¯ä»¥é›†æˆLLMï¼‰
            generation_start = time.time()
            answer = self._generate_answer(query_text, search_results)
            generation_time = time.time() - generation_start

            total_time = time.time() - start_time

            # è®°å½•æŸ¥è¯¢æ—¥å¿—
            self._log_query(
                query_text, answer, search_results,
                retrieval_time, generation_time, total_time, user
            )

            # è®°å½•æŸ¥è¯¢å®Œæˆä¿¡æ¯
            logger.info(f"âœ… çŸ¥è¯†åº“æŸ¥è¯¢å®Œæˆ:")
            logger.info(f"   â±ï¸  æ£€ç´¢è€—æ—¶: {retrieval_time:.3f}s")
            logger.info(f"   ğŸ¤– ç”Ÿæˆè€—æ—¶: {generation_time:.3f}s")
            logger.info(f"   ğŸ• æ€»è€—æ—¶: {total_time:.3f}s")
            logger.info(f"   ğŸ“Š è¿”å›ç»“æœæ•°: {len(search_results)}")

            return {
                'query': query_text,
                'answer': answer,
                'sources': search_results,
                'retrieval_time': retrieval_time,
                'generation_time': generation_time,
                'total_time': total_time
            }

        except Exception as e:
            logger.error(f"çŸ¥è¯†åº“æŸ¥è¯¢å¤±è´¥: {e}")
            raise

    def _generate_answer(self, query: str, sources: List[Dict[str, Any]]) -> str:
        """ç”Ÿæˆå›ç­”ï¼ˆç®€å•ç‰ˆæœ¬ï¼Œåç»­å¯é›†æˆLLMï¼‰"""
        if not sources:
            return "æŠ±æ­‰ï¼Œæ²¡æœ‰æ‰¾åˆ°ç›¸å…³ä¿¡æ¯ã€‚"

        # ç®€å•çš„åŸºäºæ£€ç´¢ç»“æœçš„å›ç­”ç”Ÿæˆ
        context = "\n\n".join([source['content'] for source in sources[:3]])
        return f"åŸºäºæŸ¥è¯¢ã€Œ{query}ã€æ£€ç´¢åˆ°çš„ç›¸å…³å†…å®¹ï¼š\n\n{context}"

    def _log_query(self, query: str, answer: str, sources: List[Dict[str, Any]],
                   retrieval_time: float, generation_time: float, total_time: float, user):
        """è®°å½•æŸ¥è¯¢æ—¥å¿—"""
        try:
            QueryLog.objects.create(
                knowledge_base=self.knowledge_base,
                user=user,
                query=query,
                response=answer,
                retrieved_chunks=[{
                    'content': source['content'][:200] + '...' if len(source['content']) > 200 else source['content'],
                    'metadata': source['metadata'],
                    'score': source['similarity_score']
                } for source in sources],
                similarity_scores=[source['similarity_score'] for source in sources],
                retrieval_time=retrieval_time,
                generation_time=generation_time,
                total_time=total_time
            )
        except Exception as e:
            logger.error(f"è®°å½•æŸ¥è¯¢æ—¥å¿—å¤±è´¥: {e}")

    def delete_document(self, document: Document):
        """åˆ é™¤æ–‡æ¡£"""
        try:
            # ä»å‘é‡å­˜å‚¨ä¸­åˆ é™¤
            self.vector_manager.delete_document(document)

            # åˆ é™¤æ–‡ä»¶
            if document.file:
                if os.path.exists(document.file.path):
                    os.remove(document.file.path)

            # åˆ é™¤æ•°æ®åº“è®°å½•
            document.delete()

            # æ¸…ç†å‘é‡å­˜å‚¨ç¼“å­˜ï¼ˆå› ä¸ºå†…å®¹å·²å˜åŒ–ï¼‰
            VectorStoreManager.clear_cache(self.knowledge_base.id)

            logger.info(f"æ–‡æ¡£åˆ é™¤æˆåŠŸ: {document.id}")

        except Exception as e:
            logger.error(f"åˆ é™¤æ–‡æ¡£å¤±è´¥: {e}")
            raise
